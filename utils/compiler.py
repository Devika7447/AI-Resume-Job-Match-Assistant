import os
import subprocess
import tempfile
import shutil
import requests
from jinja2 import Environment
from utils.latex_templates import TEMPLATES

# Set up Jinja2 environment with LaTeX-compatible tags
latex_env = Environment(
    block_start_string=r'\BLOCK{',
    block_end_string='}',
    variable_start_string=r'\VAR{',
    variable_end_string='}',
    comment_start_string=r'\#{',
    comment_end_string='}',
    line_statement_prefix='%%',
    line_comment_prefix='%#',
    trim_blocks=True,
    autoescape=False
)

def escape_latex_str(s: str) -> str:
    """
    Escape LaTeX special characters to prevent compilation failures.
    Uses character-by-character mapping to prevent double-escaping.
    """
    if not isinstance(s, str):
        return s
    
    res = []
    for char in s:
        if char == '\\':
            res.append('\\textbackslash{}')
        elif char in ['&', '%', '$', '_', '#', '{', '}']:
            res.append('\\' + char)
        elif char == '~':
            res.append('\\textasciitilde{}')
        elif char == '^':
            res.append('\\textasciicircum{}')
        else:
            res.append(char)
    return "".join(res)

def recursive_escape_latex(data):
    """
    Recursively traverse dictionaries and lists to escape all strings for LaTeX.
    """
    if isinstance(data, dict):
        return {k: recursive_escape_latex(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [recursive_escape_latex(item) for item in data]
    elif isinstance(data, str):
        return escape_latex_str(data)
    else:
        return data

def render_latex(template_name: str, data: dict) -> str:
    """
    Escapes variables and compiles data into the target LaTeX template using Jinja2.
    """
    template_str = TEMPLATES.get(template_name)
    if not template_str:
        raise ValueError(f"Template '{template_name}' not found.")
    
    # Escape the data structure recursively to secure latex formatting
    escaped_data = recursive_escape_latex(data)
    
    # Render with Jinja2 environment
    template = latex_env.from_string(template_str)
    return template.render(**escaped_data)

def compile_latex_local(latex_code: str) -> bytes:
    """
    Attempts to compile LaTeX code into PDF locally using pdflatex or tectonic.
    """
    # Determine local compiler
    compiler = None
    if shutil.which("pdflatex"):
        compiler = "pdflatex"
    elif shutil.which("tectonic"):
        compiler = "tectonic"
        
    if not compiler:
        raise FileNotFoundError("No local LaTeX compiler (pdflatex or tectonic) found in system PATH.")
        
    # Compile inside a temporary directory
    with tempfile.TemporaryDirectory() as tmpdir:
        tex_path = os.path.join(tmpdir, "resume.tex")
        with open(tex_path, "w", encoding="utf-8") as f:
            f.write(latex_code)
            
        if compiler == "pdflatex":
            cmd = ["pdflatex", "-interaction=nonstopmode", "resume.tex"]
        else:  # tectonic
            cmd = ["tectonic", "resume.tex"]
            
        try:
            res = subprocess.run(
                cmd,
                cwd=tmpdir,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=30
            )
            pdf_path = os.path.join(tmpdir, "resume.pdf")
            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    return f.read()
            else:
                # Retrieve errors
                log_path = os.path.join(tmpdir, "resume.log")
                error_log = ""
                if os.path.exists(log_path):
                    with open(log_path, "r", encoding="utf-8", errors="ignore") as lf:
                        error_log = lf.read()
                else:
                    error_log = res.stdout.decode("utf-8") + "\n" + res.stderr.decode("utf-8")
                raise RuntimeError(f"LaTeX Local compilation failed:\n{error_log}")
        except subprocess.TimeoutExpired:
            raise RuntimeError("Local LaTeX compilation timed out (limit: 30s).")

def compile_latex_cloud(latex_code: str) -> bytes:
    """
    Compiles LaTeX code using public online compiler APIs.
    Tries texlive.net (POST multipart) first, and falls back to latexonline.cc (GET) if it fails.
    """
    # 1. Try texlive.net POST API
    texlive_url = "https://texlive.net/cgi-bin/latexcgi"
    try:
        files = {
            'filecontents[]': ('document.tex', latex_code),
            'filename[]': (None, 'document.tex')
        }
        data = {
            'engine': 'pdflatex',
            'return': 'pdf'
        }
        response = requests.post(texlive_url, files=files, data=data, timeout=30)
        if response.status_code == 200 and response.content.startswith(b"%PDF"):
            return response.content
    except Exception as e:
        print(f"texlive.net compilation failed or timed out: {e}")
        
    # 2. Fallback to latexonline.cc GET API
    latexonline_url = "https://latexonline.cc/compile"
    try:
        response = requests.get(latexonline_url, params={'text': latex_code}, timeout=30)
        if response.status_code == 200 and response.content.startswith(b"%PDF"):
            return response.content
        else:
            err_msg = response.text[:200] if response.text else "No response body"
            raise RuntimeError(f"latexonline.cc returned status {response.status_code}: {err_msg}")
    except Exception as e:
        raise RuntimeError(f"All online compilers failed. Fallback error: {str(e)}")

def compile_latex_to_pdf(latex_code: str) -> tuple:
    """
    Tries to compile LaTeX to PDF using local pdflatex/tectonic.
    If unavailable or fails, compiles using the latexonline.cc API.
    Returns: (pdf_bytes, status_message)
    """
    local_err = None
    try:
        pdf_bytes = compile_latex_local(latex_code)
        return pdf_bytes, "Success: Resume compiled locally using system LaTeX."
    except Exception as e:
        local_err = e
        print(f"Local compilation failed: {e}")
        
    try:
        pdf_bytes = compile_latex_cloud(latex_code)
        return pdf_bytes, "Success: Resume compiled using Cloud LaTeX Compiler API."
    except Exception as cloud_err:
        total_error = (
            f"PDF Compilation Failed!\n\n"
            f"[Local Compiler Error]:\n{local_err}\n\n"
            f"[Cloud Compiler Error]:\n{cloud_err}\n\n"
            f"You can still download the .tex source file and paste it on Overleaf to download the PDF."
        )
        raise RuntimeError(total_error)

def generate_docx(resume_json: dict) -> bytes:
    """
    Generate a DOCX document representing the resume.
    """
    from docx import Document
    import io
    
    doc = Document()
    
    # Title
    doc.add_heading(resume_json.get("name", "Resume"), 0)
    
    # Contact
    contact = []
    if resume_json.get("email"): contact.append(resume_json["email"])
    if resume_json.get("phone"): contact.append(resume_json["phone"])
    if resume_json.get("linkedin"): contact.append("LinkedIn: " + resume_json["linkedin"])
    if resume_json.get("github"): contact.append("GitHub: " + resume_json["github"])
    if resume_json.get("website"): contact.append("Website: " + resume_json["website"])
    
    doc.add_paragraph(" | ".join(contact))
    
    # Summary
    if resume_json.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume_json["summary"])
        
    # Experience
    if resume_json.get("experience"):
        doc.add_heading("Experience", level=1)
        for exp in resume_json["experience"]:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('role', '')} at {exp.get('company', '')}").bold = True
            p.add_run(f" ({exp.get('dates', '')})").italic = True
            if exp.get("location"):
                p.add_run(f" - {exp['location']}")
            for pt in exp.get("bullet_points", []):
                doc.add_paragraph(pt, style="List Bullet")
                
    # Projects
    if resume_json.get("projects"):
        doc.add_heading("Projects", level=1)
        for proj in resume_json["projects"]:
            p = doc.add_paragraph()
            p.add_run(proj.get("name", "")).bold = True
            if proj.get("technologies"):
                p.add_run(f" | Technologies: {proj['technologies']}").italic = True
            if proj.get("dates"):
                p.add_run(f" ({proj['dates']})")
            for pt in proj.get("bullet_points", []):
                doc.add_paragraph(pt, style="List Bullet")
                
    # Skills
    if resume_json.get("skills"):
        doc.add_heading("Skills", level=1)
        for cat, list_skills in resume_json["skills"].items():
            p = doc.add_paragraph()
            p.add_run(f"{cat}: ").bold = True
            p.add_run(", ".join(list_skills))
            
    # Education
    if resume_json.get("education"):
        doc.add_heading("Education", level=1)
        for edu in resume_json["education"]:
            p = doc.add_paragraph()
            p.add_run(edu.get("institution", "")).bold = True
            p.add_run(f" - {edu.get('degree', '')}").italic = True
            p.add_run(f" ({edu.get('dates', '')})")
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def generate_txt(resume_json: dict) -> str:
    """
    Generate a plaintext representation of the resume.
    """
    lines = []
    lines.append(resume_json.get("name", "").upper())
    lines.append("=" * len(lines[-1]))
    contact = []
    if resume_json.get("email"): contact.append(resume_json["email"])
    if resume_json.get("phone"): contact.append(resume_json["phone"])
    if resume_json.get("linkedin"): contact.append("LinkedIn: " + resume_json["linkedin"])
    if resume_json.get("github"): contact.append("GitHub: " + resume_json["github"])
    if resume_json.get("website"): contact.append("Website: " + resume_json["website"])
    lines.append(" | ".join(contact))
    lines.append("")
    
    if resume_json.get("summary"):
        lines.append("SUMMARY")
        lines.append("-" * 7)
        lines.append(resume_json["summary"])
        lines.append("")
        
    if resume_json.get("experience"):
        lines.append("EXPERIENCE")
        lines.append("-" * 10)
        for exp in resume_json["experience"]:
            lines.append(f"{exp.get('role')} at {exp.get('company')} ({exp.get('dates')})")
            if exp.get('location'):
                lines[-1] += f" - {exp.get('location')}"
            for pt in exp.get("bullet_points", []):
                lines.append(f"  * {pt}")
            lines.append("")
            
    if resume_json.get("projects"):
        lines.append("PROJECTS")
        lines.append("-" * 8)
        for proj in resume_json["projects"]:
            lines.append(f"{proj.get('name')} | Technologies: {proj.get('technologies')} ({proj.get('dates')})")
            for pt in proj.get("bullet_points", []):
                lines.append(f"  * {pt}")
            lines.append("")
            
    if resume_json.get("skills"):
        lines.append("SKILLS")
        lines.append("-" * 6)
        for cat, list_skills in resume_json["skills"].items():
            lines.append(f"  {cat}: {', '.join(list_skills)}")
        lines.append("")
        
    if resume_json.get("education"):
        lines.append("EDUCATION")
        lines.append("-" * 9)
        for edu in resume_json["education"]:
            lines.append(f"{edu.get('institution')} | {edu.get('degree')} ({edu.get('dates')})")
            
    return "\n".join(lines)

def generate_md(resume_json: dict) -> str:
    """
    Generate a markdown representation of the resume.
    """
    lines = []
    lines.append(f"# {resume_json.get('name', '')}")
    contact = []
    if resume_json.get("email"): contact.append(f"[{resume_json['email']}](mailto:{resume_json['email']})")
    if resume_json.get("phone"): contact.append(resume_json["phone"])
    if resume_json.get("linkedin"): contact.append(f"[LinkedIn]({resume_json['linkedin']})")
    if resume_json.get("github"): contact.append(f"[GitHub]({resume_json['github']})")
    if resume_json.get("website"): contact.append(f"[Website]({resume_json['website']})")
    lines.append(" | ".join(contact))
    lines.append("")
    
    if resume_json.get("summary"):
        lines.append("## Professional Summary")
        lines.append(resume_json["summary"])
        lines.append("")
        
    if resume_json.get("experience"):
        lines.append("## Experience")
        for exp in resume_json["experience"]:
            lines.append(f"### {exp.get('role')} at **{exp.get('company')}**")
            lines.append(f"*{exp.get('dates')} | {exp.get('location')}*")
            for pt in exp.get("bullet_points", []):
                lines.append(f"- {pt}")
            lines.append("")
            
    if resume_json.get("projects"):
        lines.append("## Projects")
        for proj in resume_json["projects"]:
            lines.append(f"### {proj.get('name')}")
            lines.append(f"*Technologies: {proj.get('technologies')} | {proj.get('dates')}*")
            for pt in proj.get("bullet_points", []):
                lines.append(f"- {pt}")
            lines.append("")
            
    if resume_json.get("skills"):
        lines.append("## Skills")
        for cat, list_skills in resume_json["skills"].items():
            lines.append(f"- **{cat}**: {', '.join(list_skills)}")
        lines.append("")
        
    if resume_json.get("education"):
        lines.append("## Education")
        for edu in resume_json["education"]:
            lines.append(f"### {edu.get('institution')}")
            lines.append(f"*{edu.get('degree')} | {edu.get('dates')}*")
            
    return "\n".join(lines)
